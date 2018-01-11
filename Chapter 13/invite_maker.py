#!/usr/bin/env python3

"""Take a list of names and make a docx with custom invites for each one."""

import docx

with open('guests.txt') as f:
    names = f.readlines()
    document = docx.Document()

    for name in names:
        name = name.strip()

        document.add_paragraph('It would be a pleasure to have the company of',
                               style='Custom 1')
        document.add_paragraph(name, style='Custom 2')
        document.add_paragraph('at 11010 Memory Lane on the Evening of',
                               style='Custom 3')
        document.add_paragraph('April 1st', style='Custom 4')
        document.add_paragraph("at 7 o'clock", style='Custom 5')

        document.add_page_break()

    document.save('invites.docx')

    print("File has been created and saved as 'invites.docx'")
